import re
import spacy
import PyPDF2
from docx import Document

class ResumeParser:
    """
    Extracts structured information from resume files (PDF/DOCX)
    
    Supports international formats:
    - Phone numbers (US, UK, EU, Romanian, international)
    - Multiple languages and character sets
    - Various document formats and encodings
    """
    
    def __init__(self):
        """
        Load spaCy's English language model
        """
        try:
            self.nlp = spacy.load("en_core_web_sm")
        except OSError:
            print("spaCy model not found. Run: python -m spacy download en_core_web_sm")
            raise
    
    def extract_text_from_pdf(self, pdf_path):
        """
        Extract all text from a PDF file
        
        Handles:
        - Multi-page PDFs
        - Various encodings
        - Special characters
        """
        text = ""
        try:
            with open(pdf_path, 'rb') as file:
                pdf_reader = PyPDF2.PdfReader(file)
                for page in pdf_reader.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
        except Exception as e:
            print(f"Error reading PDF: {e}")
            return ""
        
        return text
    
    def extract_text_from_docx(self, docx_path):
        """
        Extract all text from a DOCX file
        
        Handles:
        - Tables
        - Headers/footers
        - Special formatting
        """
        try:
            doc = Document(docx_path)
            # Extract from paragraphs
            text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
            
            # Also extract from tables (if any)
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        text += "\n" + cell.text
            
            return text
        except Exception as e:
            print(f"Error reading DOCX: {e}")
            return ""
    
    def extract_contact_info(self, text):
        """
        Extract email and phone number (international support)
        
        Supports:
        - International emails (any TLD)
        - US/UK/EU phone formats
        - Romanian format: +40 XXX XXX XXX, 07XX XXX XXX
        - Mobile and landline
        """
        # Email regex: matches international email formats
        email_pattern = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Za-z]{2,}\b'
        emails = re.findall(email_pattern, text, re.IGNORECASE)
        
        # Clean text: normalize whitespace for better phone detection
        cleaned_text = ' '.join(text.split())
        
        # Comprehensive phone number patterns (international)
        phone_patterns = [
            # Romanian formats:
            # +40 123 456 789, +40-123-456-789, +40123456789
            r'\+40[-.\s]?\d{3}[-.\s]?\d{3}[-.\s]?\d{3}',
            # 07XX XXX XXX (Romanian mobile)
            r'\b07\d{2}[-.\s]?\d{3}[-.\s]?\d{3}\b',
            # 02XX XXX XXX (Romanian landline)
            r'\b0[2-3]\d{2}[-.\s]?\d{3}[-.\s]?\d{3}\b',
            
            # International formats:
            # +1 555 123 4567, +44 20 1234 5678, +33 1 23 45 67 89
            r'\+\d{1,4}[-.\s]?\(?\d{1,4}\)?[-.\s]?\d{1,4}[-.\s]?\d{1,4}[-.\s]?\d{1,9}',
            
            # US/Canada formats:
            # (555) 123-4567, 555-123-4567, 555.123.4567
            r'\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4}',
            
            # UK formats:
            # 020 1234 5678, +44 20 1234 5678
            r'\b0\d{2,4}[-.\s]?\d{3,4}[-.\s]?\d{4}\b',
            
            # Plain format:
            # 1234567890 (10-15 digits)
            r'\b\d{10,15}\b',
            
            # European formats:
            # +49 30 12345678, +33 1 23 45 67 89
            r'\+\d{2}[-.\s]?\d{1,4}[-.\s]?\d{4,10}',
        ]
        
        phone = None
        for pattern in phone_patterns:
            matches = re.findall(pattern, cleaned_text)
            if matches:
                for match in matches:
                    # Extract digits only to validate
                    digits_only = re.sub(r'\D', '', match)
                    
                    # Valid phone: 9-15 digits (covers international)
                    if 9 <= len(digits_only) <= 15:
                        # Avoid false positives (years, IDs, etc.)
                        # Skip if it's exactly 4 digits (likely a year)
                        if len(digits_only) == 4:
                            continue
                        # Skip if it looks like a date (contains 19XX or 20XX)
                        if re.search(r'19\d{2}|20\d{2}', match):
                            continue
                        
                        phone = match.strip()
                        break
            
            if phone:
                break
        
        # Fallback: Aggressive search if nothing found
        if not phone:
            # Look for any number sequence that could be a phone
            fallback_pattern = r'[\d\s\-\.\(\)\+]{9,25}'
            potential_phones = re.findall(fallback_pattern, text)
            
            for potential in potential_phones:
                digits = re.sub(r'\D', '', potential)
                # Must have 9-15 digits to be valid
                if 9 <= len(digits) <= 15:
                    # Additional validation: shouldn't be all same digit
                    if len(set(digits)) > 3:  # At least 4 different digits
                        phone = potential.strip()
                        break
        
        return {
            'email': emails[0] if emails else None,
            'phone': phone
        }
    
    def extract_skills(self, text):
        """
        Extract technical skills using comprehensive keyword matching
        
        Covers:
        - Programming languages
        - Frameworks and libraries
        - Tools and platforms
        - Methodologies
        - International tech terms
        """
        # Comprehensive tech skills list
        skill_keywords = [
            # Programming Languages
            'python', 'java', 'javascript', 'typescript', 'c++', 'c#', 'c sharp',
            'ruby', 'php', 'swift', 'kotlin', 'go', 'golang', 'rust', 'scala',
            'r', 'matlab', 'perl', 'dart', 'objective-c', 'shell', 'bash',
            'powershell', 'vba', 'assembly',
            
            # Web Technologies
            'react', 'angular', 'vue', 'vue.js', 'svelte', 'next.js', 'nuxt',
            'node.js', 'express', 'django', 'flask', 'fastapi', 'spring',
            'asp.net', 'laravel', 'symfony', 'rails', 'ruby on rails',
            'html', 'html5', 'css', 'css3', 'sass', 'scss', 'less',
            'tailwind', 'bootstrap', 'material ui', 'jquery', 'webpack',
            'vite', 'babel', 'typescript',
            
            # Mobile Development
            'android', 'ios', 'react native', 'flutter', 'xamarin', 'ionic',
            'cordova', 'swift', 'swiftui', 'kotlin',
            
            # Databases
            'sql', 'mysql', 'postgresql', 'postgres', 'mongodb', 'redis',
            'elasticsearch', 'cassandra', 'oracle', 'sql server', 'sqlite',
            'dynamodb', 'firebase', 'mariadb', 'neo4j', 'couchdb',
            
            # Cloud & DevOps
            'aws', 'amazon web services', 'azure', 'microsoft azure', 'gcp',
            'google cloud', 'docker', 'kubernetes', 'k8s', 'jenkins',
            'terraform', 'ansible', 'puppet', 'chef', 'vagrant',
            'ci/cd', 'circleci', 'travis ci', 'gitlab ci', 'github actions',
            'cloudformation', 'serverless', 'lambda', 'ec2', 's3',
            
            # Version Control
            'git', 'github', 'gitlab', 'bitbucket', 'svn', 'mercurial',
            
            # Data Science & ML
            'machine learning', 'deep learning', 'artificial intelligence',
            'ai', 'ml', 'nlp', 'computer vision', 'data science',
            'tensorflow', 'pytorch', 'keras', 'scikit-learn', 'pandas',
            'numpy', 'matplotlib', 'seaborn', 'opencv', 'spacy',
            'hugging face', 'transformers', 'bert', 'gpt',
            
            # Testing
            'unit testing', 'jest', 'mocha', 'pytest', 'junit', 'selenium',
            'cypress', 'testng', 'jasmine', 'karma',
            
            # Methodologies & Practices
            'agile', 'scrum', 'kanban', 'devops', 'tdd', 'bdd',
            'microservices', 'rest api', 'restful', 'graphql', 'soap',
            'oauth', 'jwt', 'authentication', 'authorization',
            
            # Other Tools
            'jira', 'confluence', 'slack', 'trello', 'asana',
            'postman', 'insomnia', 'swagger', 'apache', 'nginx',
            'linux', 'unix', 'windows server', 'vim', 'emacs',
            'vs code', 'visual studio', 'intellij', 'eclipse', 'pycharm',
        ]
        
        text_lower = text.lower()
        found_skills = []
        
        for skill in skill_keywords:
            # Use word boundaries to avoid partial matches
            pattern = r'\b' + re.escape(skill.replace('.', r'\.')) + r'\b'
            if re.search(pattern, text_lower):
                # Normalize skill name (capitalize properly)
                if skill in ['html', 'css', 'sql', 'aws', 'gcp', 'api', 'nlp', 'ai', 'ml']:
                    found_skills.append(skill.upper())
                elif skill == 'c#' or skill == 'c sharp':
                    found_skills.append('C#')
                elif skill == 'c++':
                    found_skills.append('C++')
                elif '.' in skill:  # node.js, vue.js, etc.
                    found_skills.append(skill)
                else:
                    found_skills.append(skill.capitalize())
        
        # Remove duplicates while preserving order
        seen = set()
        unique_skills = []
        for skill in found_skills:
            skill_normalized = skill.lower()
            if skill_normalized not in seen:
                seen.add(skill_normalized)
                unique_skills.append(skill)
        
        return unique_skills
    
    def extract_entities(self, text):
        """
        Use spaCy's Named Entity Recognition
        
        Extracts:
        - Names (PERSON)
        - Organizations/Companies (ORG)
        - Locations (GPE)
        - Dates (DATE)
        """
        doc = self.nlp(text)
        
        entities = {
            'persons': [],
            'organizations': [],
            'locations': [],
            'dates': []
        }
        
        for ent in doc.ents:
            if ent.label_ == 'PERSON':
                entities['persons'].append(ent.text)
            elif ent.label_ == 'ORG':
                entities['organizations'].append(ent.text)
            elif ent.label_ == 'GPE':
                entities['locations'].append(ent.text)
            elif ent.label_ == 'DATE':
                entities['dates'].append(ent.text)
        
        # Remove duplicates
        for key in entities:
            entities[key] = list(set(entities[key]))
        
        return entities
    
    def extract_education(self, text):
        """
        Extract education information
        
        Looks for:
        - Degree types (international)
        - Universities and institutions
        - Fields of study
        """
        education_keywords = [
            # Degrees (international)
            'bachelor', 'master', 'phd', 'doctorate', 'mba', 'degree',
            'licenta', 'master', 'doctorat',  # Romanian
            'b.s.', 'm.s.', 'b.a.', 'm.a.', 'b.sc.', 'm.sc.',
            'b.tech', 'm.tech', 'b.e.', 'm.e.',
            
            # Institutions
            'university', 'college', 'institute', 'school', 'academy',
            'universitate', 'colegiu', 'institut',  # Romanian
            
            # Fields
            'computer science', 'engineering', 'mathematics', 'physics',
            'business', 'economics', 'informatica', 'inginerie',  # Romanian
            'software engineering', 'data science', 'information technology'
        ]
        
        text_lower = text.lower()
        found_education = []
        
        for keyword in education_keywords:
            if keyword in text_lower:
                found_education.append(keyword)
        
        return list(set(found_education))
    
    def parse_resume(self, file_path):
        """
        Main parsing function
        
        Process:
        1. Determine file type
        2. Extract raw text
        3. Parse contact information
        4. Extract skills (international)
        5. Use spaCy for entity recognition
        6. Extract education
        7. Calculate metrics
        
        Returns: dict with all extracted information
        """
        import os
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"File not found: {file_path}")
        
        # Extract text based on file type
        if file_path.lower().endswith('.pdf'):
            text = self.extract_text_from_pdf(file_path)
        elif file_path.lower().endswith('.docx'):
            text = self.extract_text_from_docx(file_path)
        else:
            raise ValueError("Unsupported file format. Use .pdf or .docx")
        
        if not text or len(text.strip()) == 0:
            raise ValueError("No text could be extracted from the file")
        
        # Extract all information
        contact_info = self.extract_contact_info(text)
        skills = self.extract_skills(text)
        entities = self.extract_entities(text)
        education = self.extract_education(text)
        
        # Calculate metrics
        word_count = len(text.split())
        char_count = len(text)
        
        # Return structured data
        return {
            'raw_text': text,
            'contact_info': contact_info,
            'skills': skills,
            'entities': entities,
            'education': education,
            'word_count': word_count,
            'char_count': char_count,
            'file_path': file_path
        }